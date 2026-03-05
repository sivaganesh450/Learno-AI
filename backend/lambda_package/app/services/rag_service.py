"""
Agentic RAG Service for Summarizer Agent
Implements Conversational Agentic RAG Architecture:
1. History Aware Query Rephraser - Rephrases follow-up questions into standalone queries
2. Router Agent - Decides whether to use documents, web search, or both
3. In-Memory Document Store - Stores document chunks with keyword-based retrieval (Lambda-compatible)
4. Tavily Web Search - For external/current information
5. RAG Pipeline - Retrieves context and generates answers
6. Amazon Textract - For extracting text from images
"""

import os
import io
import tempfile
from typing import List, Dict, Optional
from pathlib import Path

# Document loaders
from pypdf import PdfReader
from docx import Document as DocxDocument

# Image processing
from PIL import Image

# AWS services
import boto3

# LangChain components
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.messages import HumanMessage, AIMessage, SystemMessage
from langchain_community.chat_message_histories import ChatMessageHistory
from langchain_aws import ChatBedrockConverse

# ChromaDB removed — using lightweight in-memory document store for Lambda compatibility

# Tavily for web search
from tavily import TavilyClient


class RAGService:
    """Agentic RAG Service for document-based Q&A with web search capabilities"""
    
    def __init__(self):
        self.name = "Agentic RAG Summarizer System"
        
        # Initialize Tavily for web search
        TAVILY_API_KEY = os.environ.get("TAVILY_API_KEY", "")
        try:
            self.tavily = TavilyClient(api_key=TAVILY_API_KEY) if TAVILY_API_KEY else None
            if self.tavily:
                print("Tavily client initialized for Agentic RAG")
            else:
                print("Tavily not configured for RAG (TAVILY_API_KEY not set)")
        except Exception as e:
            print(f"Failed to initialize Tavily: {e}")
            self.tavily = None
        
        # In-memory document store (replaces ChromaDB for Lambda compatibility)
        # Stores chunked documents per session for retrieval
        self.document_chunks: Dict[str, List[Dict]] = {}
        print("In-memory document store initialized")
        
        # Initialize LLM (Amazon Bedrock - Llama 3.3 70B)
        try:
            self.llm = ChatBedrockConverse(
                model=os.getenv('BEDROCK_MODEL_ID', 'us.meta.llama3-3-70b-instruct-v1:0'),
                region_name=os.getenv('AWS_REGION', 'us-east-1'),
                max_tokens=2048,
                temperature=0.5,
            )
            print("Bedrock LLM initialized for RAG (Llama 3.3 70B)")
        except Exception as e:
            print(f"Failed to initialize Bedrock for RAG: {e}")
            self.llm = None
        
        # Text splitter for chunking documents
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        
        # Session-based conversation history (k=3 turns like ConversationBufferWindowMemory)
        self.session_histories: Dict[str, ChatMessageHistory] = {}
        self.k = 3  # Keep last k conversation turns
        
        # Track uploaded documents per session
        self.session_documents: Dict[str, List[str]] = {}
        
        # Keywords that indicate user wants external/web information
        self.web_search_indicators = [
            'latest', 'recent', 'current', 'today', 'news', '2024', '2025', '2026',
            'search online', 'search the web', 'look up', 'find online',
            'what is happening', 'trending', 'update', 'external'
        ]
    
    def _needs_web_search(self, query: str) -> bool:
        """Determine if query needs web search based on keywords"""
        query_lower = query.lower()
        return any(indicator in query_lower for indicator in self.web_search_indicators)
    
    def _search_web(self, query: str) -> str:
        """Search the web using Tavily API"""
        if not self.tavily:
            return ""
        
        try:
            print(f"[Agentic RAG] Searching web for: {query}")
            response = self.tavily.search(
                query=query,
                search_depth="advanced",
                max_results=5,
                include_answer=True,
                include_raw_content=False
            )
            
            results = []
            
            # Include Tavily's AI-generated answer if available
            if response.get('answer'):
                results.append(f"**Web Summary:** {response['answer']}")
            
            # Include top search results
            for i, result in enumerate(response.get('results', [])[:3], 1):
                title = result.get('title', 'No title')
                content = result.get('content', '')[:300]
                url = result.get('url', '')
                results.append(f"[{i}] **{title}**\n{content}...\n🔗 {url}")
            
            web_context = "\n\n".join(results)
            print(f"[Agentic RAG] Found {len(response.get('results', []))} web results")
            return web_context
            
        except Exception as e:
            print(f"[Agentic RAG] Web search error: {e}")
            return ""
    
    def get_session_history(self, session_id: str) -> ChatMessageHistory:
        """Get or create ChatMessageHistory for a session"""
        if session_id not in self.session_histories:
            self.session_histories[session_id] = ChatMessageHistory()
        return self.session_histories[session_id]
    
    def get_windowed_history(self, session_id: str) -> List:
        """Get last k conversation turns"""
        history = self.get_session_history(session_id)
        messages = history.messages
        return messages[-(self.k * 2):] if len(messages) > self.k * 2 else messages
    
    def save_context(self, session_id: str, user_input: str, ai_output: str):
        """Save conversation context"""
        history = self.get_session_history(session_id)
        history.add_user_message(user_input)
        history.add_ai_message(ai_output)
        print(f"[RAG Memory] Saved context for session {session_id}. Total messages: {len(history.messages)}")
    
    def clear_session(self, session_id: str):
        """Clear session history and documents"""
        if session_id in self.session_histories:
            del self.session_histories[session_id]
        if session_id in self.session_documents:
            del self.session_documents[session_id]
        if session_id in self.document_chunks:
            del self.document_chunks[session_id]
            print(f"[RAG] Cleared document chunks for session {session_id}")
    
    def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file using multiple strategies."""
        # Strategy 1: pypdf standard extraction
        text = ""
        try:
            reader = PdfReader(file_path)
            print(f"[RAG PDF] Opened PDF with {len(reader.pages)} pages")
            
            for i, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text += page_text + "\n"
                        print(f"[RAG PDF] Page {i+1}: extracted {len(page_text)} chars")
                    else:
                        print(f"[RAG PDF] Page {i+1}: no text extracted")
                except Exception as page_err:
                    print(f"[RAG PDF] Page {i+1} error: {page_err}")
            
            if text.strip():
                print(f"[RAG PDF] pypdf total: {len(text)} chars")
                return text
            else:
                print(f"[RAG PDF] pypdf extracted 0 chars of actual text")
        except Exception as e:
            print(f"[RAG PDF] pypdf failed: {type(e).__name__}: {e}")
        
        # Strategy 2: Extract raw text from PDF bytes
        print(f"[RAG PDF] Trying raw text extraction from PDF bytes")
        text = self._extract_raw_text_from_pdf(file_path)
        if text.strip():
            print(f"[RAG PDF] Raw extraction got {len(text)} chars")
            return text
        
        # Strategy 3: Try reading as plain text (some PDFs are actually text)
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                raw = f.read()
            # Filter to printable text content
            import re
            # Extract text between BT...ET text blocks in PDF
            text_blocks = re.findall(r'\(([^)]+)\)', raw)
            if text_blocks:
                filtered = ' '.join(t for t in text_blocks if len(t) > 3 and any(c.isalpha() for c in t))
                if len(filtered) > 100:
                    print(f"[RAG PDF] Raw regex extraction got {len(filtered)} chars")
                    return filtered
        except Exception as e:
            print(f"[RAG PDF] Raw text read failed: {e}")
        
        print(f"[RAG PDF] All extraction strategies failed")
        return ""
    
    def _extract_raw_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF using pypdf with different layout modes."""
        try:
            reader = PdfReader(file_path)
            text = ""
            
            for page in reader.pages:
                # Try with layout extraction
                try:
                    # Try extracting with visitor pattern to get all text
                    parts = []
                    def visitor_text(text_content, cm, tm, fontDict, fontSize):
                        if text_content and text_content.strip():
                            parts.append(text_content)
                    
                    page.extract_text(visitor_text=visitor_text)
                    if parts:
                        text += ' '.join(parts) + "\n"
                except Exception:
                    pass
            
            return text
        except Exception as e:
            print(f"[RAG PDF] Raw extraction error: {e}")
            return ""
    
    def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        text = ""
        try:
            doc = DocxDocument(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
        except Exception as e:
            print(f"Error extracting DOCX: {e}")
        return text
    
    def _extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            print(f"Error reading TXT: {e}")
            return ""
    
    def _extract_text_from_image(self, file_path: str) -> str:
        """Extract text from image using Amazon Textract"""
        try:
            textract = boto3.client('textract', region_name=os.getenv('AWS_REGION', 'us-east-1'))
            
            with open(file_path, 'rb') as f:
                image_bytes = f.read()
            
            response = textract.detect_document_text(
                Document={'Bytes': image_bytes}
            )
            
            text = ' '.join([
                block['Text'] for block in response['Blocks']
                if block['BlockType'] == 'LINE'
            ])
            
            if text.strip():
                print(f"[Textract] Extracted text from image: {len(text)} chars")
                return text
            
            return "[No text detected in image]"
            
        except Exception as e:
            print(f"[Textract] Error extracting text from image: {e}")
            return f"[Error extracting text from image: {str(e)}]"

    async def process_document(self, file_content: bytes, filename: str, session_id: str) -> Dict:
        """Process uploaded document and add to in-memory store"""
        print(f"[RAG] Processing document: {filename} ({len(file_content)} bytes) for session: {session_id}")
        
        # Save to temp file
        suffix = Path(filename).suffix.lower()
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix, dir='/tmp') as tmp:
                tmp.write(file_content)
                tmp_path = tmp.name
            print(f"[RAG] Saved temp file: {tmp_path} ({os.path.getsize(tmp_path)} bytes)")
        except Exception as e:
            print(f"[RAG] Failed to write temp file: {e}")
            return {"success": False, "error": f"Failed to save uploaded file: {str(e)}"}
        
        try:
            # Extract text based on file type
            if suffix == '.pdf':
                text = self._extract_text_from_pdf(tmp_path)
            elif suffix in ['.docx', '.doc']:
                text = self._extract_text_from_docx(tmp_path)
            elif suffix in ['.txt', '.md', '.py', '.js', '.json', '.csv']:
                text = self._extract_text_from_txt(tmp_path)
            elif suffix in ['.png', '.jpg', '.jpeg', '.bmp', '.gif', '.tiff', '.webp']:
                # Use Amazon Textract for image text extraction
                print(f"[RAG] Using Textract for image: {filename}")
                text = self._extract_text_from_image(tmp_path)
            else:
                return {"success": False, "error": f"Unsupported file type: {suffix}"}
            
            if not text.strip():
                print(f"[RAG] No text extracted from {filename} (suffix={suffix})")
                return {"success": False, "error": f"Could not extract text from document. The file may be image-based/scanned or empty. Try uploading a text-based PDF or a .txt/.docx file instead."}
            
            # Split text into chunks
            chunks = self.text_splitter.split_text(text)
            print(f"[RAG] Split document into {len(chunks)} chunks")
            
            # Create document ID
            doc_id = f"{session_id}_{filename}"
            
            # Store chunks in memory
            if session_id not in self.document_chunks:
                self.document_chunks[session_id] = []
            
            for i, chunk in enumerate(chunks):
                self.document_chunks[session_id].append({
                    "id": f"{doc_id}_{i}",
                    "text": chunk,
                    "filename": filename,
                    "source": doc_id,
                    "chunk_index": i
                })
            
            print(f"[RAG] Added {len(chunks)} chunks to in-memory store")
            
            # Track document for session
            if session_id not in self.session_documents:
                self.session_documents[session_id] = []
            self.session_documents[session_id].append(doc_id)
            
            return {
                "success": True,
                "filename": filename,
                "chunks": len(chunks),
                "characters": len(text)
            }
            
        finally:
            # Clean up temp file
            os.unlink(tmp_path)
    
    def _format_history_for_rephraser(self, history_messages: List) -> str:
        """Format conversation history for the rephraser prompt"""
        if not history_messages:
            return "No previous conversation."
        
        formatted = []
        for msg in history_messages:
            if isinstance(msg, HumanMessage):
                formatted.append(f"Human: {msg.content}")
            elif isinstance(msg, AIMessage):
                # Truncate long AI responses
                content = msg.content[:500] + "..." if len(msg.content) > 500 else msg.content
                formatted.append(f"Assistant: {content}")
        
        return "\n".join(formatted)
    
    async def _rephrase_query(self, query: str, history_messages: List) -> str:
        """
        History Aware Query Rephraser
        Rephrases follow-up questions into standalone queries using conversation context
        """
        if not history_messages or not self.llm:
            return query
        
        history_text = self._format_history_for_rephraser(history_messages)
        
        rephraser_prompt = f"""Given a query and historical conversation, rephrase the query if it refers to something from the historical conversation into a standalone question that doesn't need the history to understand.

If the query is already a standalone question, return it as-is.

Conversation History:
{history_text}

Current Query: {query}

Rephrased Query (standalone question):"""
        
        try:
            response = await self.llm.ainvoke([HumanMessage(content=rephraser_prompt)])
            rephrased = response.content.strip()
            print(f"[RAG Rephraser] Original: '{query}' → Rephrased: '{rephrased}'")
            return rephrased
        except Exception as e:
            print(f"[RAG Rephraser] Error: {e}, using original query")
            return query
    
    def _retrieve_context(self, query: str, session_id: str, k: int = 6) -> List[Document]:
        """Retrieve relevant document chunks from in-memory store using keyword matching"""
        chunks = self.document_chunks.get(session_id, [])
        if not chunks:
            return []
        
        try:
            # Simple keyword-based relevance scoring
            query_words = set(query.lower().split())
            # Remove common stop words
            stop_words = {'the', 'a', 'an', 'is', 'are', 'was', 'were', 'in', 'on', 'at', 'to', 
                         'for', 'of', 'with', 'and', 'or', 'not', 'it', 'this', 'that', 'what',
                         'how', 'why', 'when', 'where', 'which', 'who', 'do', 'does', 'did',
                         'can', 'could', 'will', 'would', 'should', 'may', 'might', 'be', 'been',
                         'have', 'has', 'had', 'from', 'by', 'about', 'as', 'into', 'its', 'my'}
            query_keywords = query_words - stop_words
            
            scored_chunks = []
            for chunk in chunks:
                chunk_text_lower = chunk["text"].lower()
                # Score = number of query keywords found in chunk
                score = sum(1 for word in query_keywords if word in chunk_text_lower)
                scored_chunks.append((score, chunk))
            
            # Sort by score descending, take top k
            scored_chunks.sort(key=lambda x: x[0], reverse=True)
            
            docs = []
            for score, chunk in scored_chunks[:k]:
                docs.append(Document(
                    page_content=chunk["text"],
                    metadata={
                        "filename": chunk["filename"],
                        "source": chunk["source"],
                        "chunk_index": chunk["chunk_index"]
                    }
                ))
                print(f"[RAG Retrieval] Chunk '{chunk['filename']}' idx={chunk['chunk_index']} score={score}")
            
            print(f"[RAG Retrieval] Found {len(docs)} chunks for query: '{query[:50]}...'")
            return docs
        except Exception as e:
            print(f"[RAG Retrieval] Error: {e}")
            return []
    
    def _format_context(self, documents: List) -> str:
        """Format retrieved documents into context string"""
        if not documents:
            return "No relevant documents found."
        
        context_parts = []
        for i, doc in enumerate(documents, 1):
            if hasattr(doc, 'metadata'):
                source = doc.metadata.get('filename', 'Unknown')
                content = doc.page_content
            else:
                source = 'Unknown'
                content = str(doc)
            context_parts.append(f"[Chunk {i} from {source}]\n{content}")
        
        return "\n\n".join(context_parts)
    
    async def answer_stream(self, question: str, session_id: str = "default"):
        """
        Answer question using Agentic RAG with streaming
        
        Architecture:
        1. Get conversation history
        2. Rephrase query using history (History Aware Query Rephraser)
        3. Router: Decide if web search is needed
        4. Retrieve relevant context from document store (if documents uploaded)
        5. Search web with Tavily (if needed)
        6. Generate answer using combined context
        """
        if not self.llm:
            yield "Error: LLM not initialized. Please ensure AWS Bedrock is configured."
            return
        
        # Step 1: Get conversation history
        history_messages = self.get_windowed_history(session_id)
        print(f"[Agentic RAG] Session: {session_id}, History messages: {len(history_messages)}")
        
        # Step 2: Rephrase query using history (History Aware Query Rephraser)
        rephrased_query = await self._rephrase_query(question, history_messages)
        
        # Step 3: Router - Check if web search is needed
        needs_web = self._needs_web_search(question)
        has_documents = self.get_document_count(session_id) > 0
        
        # Step 4: Retrieve relevant context from vector store (if documents exist)
        doc_context = ""
        has_relevant_docs = False
        if has_documents:
            retrieved_docs = self._retrieve_context(rephrased_query, session_id)
            doc_context = self._format_context(retrieved_docs)
            # Check if we actually got relevant content
            has_relevant_docs = doc_context and doc_context != "No relevant documents found."
        
        # Step 5: Search web if needed OR if documents don't have relevant info
        web_context = ""
        if needs_web or (not has_relevant_docs):
            web_context = self._search_web(rephrased_query)
        
        # Combine contexts
        combined_context = ""
        sources_used = []
        
        if doc_context and doc_context != "No relevant documents found.":
            combined_context += f"📄 **From Uploaded Documents:**\n{doc_context}\n\n"
            sources_used.append("documents")
        
        if web_context:
            combined_context += f"🌐 **From Web Search:**\n{web_context}\n\n"
            sources_used.append("web")
        
        if not combined_context:
            combined_context = "No relevant information found in documents or web search."
        
        # Format history for RAG prompt
        history_text = self._format_history_for_rephraser(history_messages)
        
        print(f"[Agentic RAG] Sources used: {sources_used}, Needs web: {needs_web}, Has docs: {has_documents}")
        
        # Step 6: Generate answer using Agentic RAG prompt
        rag_system_prompt = """You are a helpful AI assistant with access to both uploaded documents and web search results.

IMPORTANT RULES:
1. ALWAYS use the provided context to answer - this includes BOTH document context AND web search results
2. If web search results are provided, USE THEM to answer the question directly
3. Do NOT ask for permission to search the web - the search has already been done for you
4. If using documents, reference the source when possible
5. If context includes web results, provide the answer from those results
6. Consider the conversation history for context
7. Use clear, concise language with markdown formatting
8. Do NOT make up information - only use what's provided in the context"""

        rag_user_prompt = f"""Query: {question}

Retrieved Context:
{combined_context}

Conversation History:
{history_text}

Answer the query using the context above. If web search results are provided, use them to give a complete answer. Do NOT ask user permission to search - just use the provided results."""

        messages = [
            SystemMessage(content=rag_system_prompt),
            HumanMessage(content=rag_user_prompt)
        ]
        
        # Stream the response
        full_response = ""
        try:
            async for chunk in self.llm.astream(messages):
                if hasattr(chunk, 'content') and chunk.content:
                    full_response += chunk.content
                    yield chunk.content
            
            # Save context after streaming completes
            if full_response:
                self.save_context(session_id, question, full_response)
                
        except Exception as e:
            print(f"[Agentic RAG] Streaming error: {e}")
            yield f"Error generating response: {str(e)}"
    
    async def answer(self, question: str, session_id: str = "default") -> str:
        """Non-streaming version of answer"""
        full_response = ""
        async for chunk in self.answer_stream(question, session_id):
            full_response += chunk
        return full_response
    
    def get_session_documents(self, session_id: str) -> List[str]:
        """Get list of uploaded documents for a session"""
        return self.session_documents.get(session_id, [])
    
    def get_document_count(self, session_id: str) -> int:
        """Get count of documents for a session"""
        return len(self.session_documents.get(session_id, []))
    
    def get_chunk_count(self, session_id: str) -> int:
        """Get count of document chunks for a session"""
        return len(self.document_chunks.get(session_id, []))


# Global instance
rag_service = RAGService()
