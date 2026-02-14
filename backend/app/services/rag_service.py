"""
Agentic RAG Service for Summarizer Agent
Implements Conversational Agentic RAG Architecture:
1. History Aware Query Rephraser - Rephrases follow-up questions into standalone queries
2. Router Agent - Decides whether to use documents, web search, or both
3. Vector DB (ChromaDB) - Stores document embeddings for retrieval
4. Tavily Web Search - For external/current information
5. RAG Pipeline - Retrieves context and generates answers
6. Tesseract OCR - For extracting text from images
"""

import os
import io
import tempfile
from typing import List, Dict, Optional
from pathlib import Path

# Document loaders
from pypdf import PdfReader
from docx import Document as DocxDocument

# Image processing
from PIL import Image

# Tesseract for OCR
try:
    import pytesseract
    pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
    TESSERACT_AVAILABLE = True
    print("Tesseract OCR initialized for image text extraction")
except:
    TESSERACT_AVAILABLE = False
    print("Tesseract not available for image OCR")

# LangChain components
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.messages import HumanMessage, AIMessage, SystemMessage
from langchain_community.chat_message_histories import ChatMessageHistory
# from langchain_ollama import ChatOllama # Removed Ollama
from google import genai

# ChromaDB with default embeddings
import chromadb
from chromadb.config import Settings
from chromadb.utils import embedding_functions

# Tavily for web search
from tavily import TavilyClient


class RAGService:
    """Agentic RAG Service for document-based Q&A with web search capabilities"""
    
    def __init__(self):
        self.name = "Agentic RAG Summarizer System"
        
        # Initialize Tavily for web search
        TAVILY_API_KEY = "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"
        try:
            self.tavily = TavilyClient(api_key=TAVILY_API_KEY)
            print("Tavily client initialized for Agentic RAG")
        except Exception as e:
            print(f"Failed to initialize Tavily: {e}")
            self.tavily = None
        
        print("Initializing ChromaDB...")
        
        # Use ephemeral client (in-memory) to avoid persistence issues
        # Documents will be re-uploaded each session
        try:
            self.chroma_client = chromadb.Client()
            
            # Use ChromaDB's default embedding function
            self.embedding_function = embedding_functions.DefaultEmbeddingFunction()
            
            # Get or create collection with embedding function
            self.collection = self.chroma_client.get_or_create_collection(
                name="lerno_documents",
                embedding_function=self.embedding_function,
                metadata={"hnsw:space": "cosine"}
            )
            print("ChromaDB initialized (in-memory)")
        except Exception as e:
            print(f"ChromaDB initialization error: {e}")
            self.collection = None
        
        # Initialize Gemini Client
        try:
            self.client = genai.Client(api_key=os.environ.get("GOOGLE_API_KEY"))
            print("Gemini client initialized for RAG")
        except Exception as e:
            print(f"Failed to initialize Gemini: {e}")
            self.client = None
        
        # Text splitter for chunking documents
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        
        # Session-based conversation history (k=3 turns like ConversationBufferWindowMemory)
        self.session_histories: Dict[str, ChatMessageHistory] = {}
        self.k = 3  # Keep last k conversation turns
        
        # Track uploaded documents per session
        self.session_documents: Dict[str, List[str]] = {}
        
        # Keywords that indicate user wants external/web information
        self.web_search_indicators = [
            'latest', 'recent', 'current', 'today', 'news', '2024', '2025', '2026',
            'search online', 'search the web', 'look up', 'find online',
            'what is happening', 'trending', 'update', 'external'
        ]
    
    def _needs_web_search(self, query: str) -> bool:
        """Determine if query needs web search based on keywords"""
        query_lower = query.lower()
        return any(indicator in query_lower for indicator in self.web_search_indicators)
    
    def _search_web(self, query: str) -> str:
        """Search the web using Tavily API"""
        if not self.tavily:
            return ""
        
        try:
            print(f"[Agentic RAG] Searching web for: {query}")
            response = self.tavily.search(
                query=query,
                search_depth="advanced",
                max_results=5,
                include_answer=True,
                include_raw_content=False
            )
            
            results = []
            
            # Include Tavily's AI-generated answer if available
            if response.get('answer'):
                results.append(f"**Web Summary:** {response['answer']}")
            
            # Include top search results
            for i, result in enumerate(response.get('results', [])[:3], 1):
                title = result.get('title', 'No title')
                content = result.get('content', '')[:300]
                url = result.get('url', '')
                results.append(f"[{i}] **{title}**\n{content}...\n🔗 {url}")
            
            web_context = "\n\n".join(results)
            print(f"[Agentic RAG] Found {len(response.get('results', []))} web results")
            return web_context
            
        except Exception as e:
            print(f"[Agentic RAG] Web search error: {e}")
            return ""
    
    def get_session_history(self, session_id: str) -> ChatMessageHistory:
        """Get or create ChatMessageHistory for a session"""
        if session_id not in self.session_histories:
            self.session_histories[session_id] = ChatMessageHistory()
        return self.session_histories[session_id]
    
    def get_windowed_history(self, session_id: str) -> List:
        """Get last k conversation turns"""
        history = self.get_session_history(session_id)
        messages = history.messages
        return messages[-(self.k * 2):] if len(messages) > self.k * 2 else messages
    
    def save_context(self, session_id: str, user_input: str, ai_output: str):
        """Save conversation context"""
        history = self.get_session_history(session_id)
        history.add_user_message(user_input)
        history.add_ai_message(ai_output)
        print(f"[RAG Memory] Saved context for session {session_id}. Total messages: {len(history.messages)}")
    
    def clear_session(self, session_id: str):
        """Clear session history and documents"""
        if session_id in self.session_histories:
            del self.session_histories[session_id]
        if session_id in self.session_documents:
            # Delete documents from ChromaDB
            for doc_id in self.session_documents[session_id]:
                try:
                    self.collection.delete(where={"source": doc_id})
                except Exception as e:
                    print(f"Error deleting document {doc_id}: {e}")
            del self.session_documents[session_id]
    
    def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            reader = PdfReader(file_path)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        except Exception as e:
            print(f"Error extracting PDF: {e}")
        return text
    
    def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        text = ""
        try:
            doc = DocxDocument(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
        except Exception as e:
            print(f"Error extracting DOCX: {e}")
        return text
    
    def _extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            print(f"Error reading TXT: {e}")
            return ""
    
    def _extract_text_from_image(self, file_path: str) -> str:
        """Extract text from image using Tesseract OCR"""
        if not TESSERACT_AVAILABLE:
            return "[Image uploaded but Tesseract OCR not available. Install Tesseract-OCR]"
        
        try:
            # Open and preprocess image
            image = Image.open(file_path)
            
            # Convert to grayscale for better OCR
            if image.mode != 'L':
                image = image.convert('L')
            
            # Try different PSM modes for best result
            best_text = ""
            for psm in [6, 4, 3, 11]:
                config = f'--oem 3 --psm {psm} -c preserve_interword_spaces=1'
                text = pytesseract.image_to_string(image, config=config)
                if text and len(text.strip()) > len(best_text):
                    best_text = text.strip()
            
            if best_text:
                print(f"[Tesseract] Extracted text from image: {len(best_text)} chars")
                return best_text
            
            return "[No text detected in image]"
            
        except Exception as e:
            print(f"[Tesseract] Error extracting text from image: {e}")
            return f"[Error extracting text from image: {str(e)}]"

    async def process_document(self, file_content: bytes, filename: str, session_id: str) -> Dict:
        """Process uploaded document and add to vector store"""
        if not self.collection:
            return {"success": False, "error": "ChromaDB not initialized."}
        
        print(f"[RAG] Processing document: {filename} for session: {session_id}")
        
        # Save to temp file
        suffix = Path(filename).suffix.lower()
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(file_content)
            tmp_path = tmp.name
        
        try:
            # Extract text based on file type
            if suffix == '.pdf':
                text = self._extract_text_from_pdf(tmp_path)
            elif suffix in ['.docx', '.doc']:
                text = self._extract_text_from_docx(tmp_path)
            elif suffix in ['.txt', '.md', '.py', '.js', '.json', '.csv']:
                text = self._extract_text_from_txt(tmp_path)
            elif suffix in ['.png', '.jpg', '.jpeg', '.bmp', '.gif', '.tiff', '.webp']:
                # Use TrOCR for image text extraction
                print(f"[RAG] Using TrOCR for image: {filename}")
                text = self._extract_text_from_image(tmp_path)
                text = self._extract_text_from_txt(tmp_path)
            else:
                return {"success": False, "error": f"Unsupported file type: {suffix}"}
            
            if not text.strip():
                return {"success": False, "error": "Could not extract text from document"}
            
            # Split text into chunks
            chunks = self.text_splitter.split_text(text)
            print(f"[RAG] Split document into {len(chunks)} chunks")
            
            # Create document ID and prepare for ChromaDB
            doc_id = f"{session_id}_{filename}"
            
            # Add to ChromaDB collection directly
            ids = [f"{doc_id}_{i}" for i in range(len(chunks))]
            metadatas = [
                {
                    "source": doc_id,
                    "filename": filename,
                    "session_id": session_id,
                    "chunk_index": i
                }
                for i in range(len(chunks))
            ]
            
            self.collection.add(
                ids=ids,
                documents=chunks,
                metadatas=metadatas
            )
            print(f"[RAG] Added {len(chunks)} chunks to ChromaDB")
            
            # Track document for session
            if session_id not in self.session_documents:
                self.session_documents[session_id] = []
            self.session_documents[session_id].append(doc_id)
            
            return {
                "success": True,
                "filename": filename,
                "chunks": len(chunks),
                "characters": len(text)
            }
            
        finally:
            # Clean up temp file
            os.unlink(tmp_path)
    
    def _format_history_for_rephraser(self, history_messages: List) -> str:
        """Format conversation history for the rephraser prompt"""
        if not history_messages:
            return "No previous conversation."
        
        formatted = []
        for msg in history_messages:
            if isinstance(msg, HumanMessage):
                formatted.append(f"Human: {msg.content}")
            elif isinstance(msg, AIMessage):
                # Truncate long AI responses
                content = msg.content[:500] + "..." if len(msg.content) > 500 else msg.content
                formatted.append(f"Assistant: {content}")
        
        return "\n".join(formatted)
    
    async def _rephrase_query(self, query: str, history_messages: List) -> str:
        """
        History Aware Query Rephraser
        Rephrases follow-up questions into standalone queries using conversation context
        """
        if not history_messages or not self.client:
            return query
        
        history_text = self._format_history_for_rephraser(history_messages)
        
        rephraser_prompt = f"""Given a query and historical conversation, rephrase the query if it refers to something from the historical conversation into a standalone question that doesn't need the history to understand.

If the query is already a standalone question, return it as-is.

Conversation History:
{history_text}

Current Query: {query}

Rephrased Query (standalone question):"""
        
        try:
            response = self.client.models.generate_content(
                model="gemini-2.0-flash-lite-001",
                contents=rephraser_prompt
            )
            rephrased = response.text.strip()
            print(f"[RAG Rephraser] Original: '{query}' → Rephrased: '{rephrased}'")
            return rephrased
        except Exception as e:
            print(f"[RAG Rephraser] Error: {e}, using original query")
            return query
    
    def _retrieve_context(self, query: str, session_id: str, k: int = 4) -> List[Document]:
        """Retrieve relevant document chunks from vector store"""
        if not self.collection:
            return []
        
        try:
            # Search with session filter using ChromaDB query
            # Include distances to filter by relevance
            results = self.collection.query(
                query_texts=[query],
                n_results=k,
                where={"session_id": session_id},
                include=["documents", "metadatas", "distances"]
            )
            
            # Convert to list of Document objects, filtering by relevance score
            # ChromaDB uses L2 distance - lower is better. Threshold of 1.0 filters irrelevant results
            RELEVANCE_THRESHOLD = 1.0
            docs = []
            if results and results['documents'] and len(results['documents']) > 0:
                for i, doc_text in enumerate(results['documents'][0]):
                    distance = results['distances'][0][i] if results.get('distances') else 0
                    # Only include documents with good relevance score
                    if distance < RELEVANCE_THRESHOLD:
                        metadata = results['metadatas'][0][i] if results['metadatas'] else {}
                        docs.append(Document(page_content=doc_text, metadata=metadata))
                        print(f"[RAG Retrieval] Chunk {i+1} distance: {distance:.3f} - RELEVANT")
                    else:
                        print(f"[RAG Retrieval] Chunk {i+1} distance: {distance:.3f} - FILTERED OUT")
            
            print(f"[RAG Retrieval] Found {len(docs)} relevant chunks for query: '{query[:50]}...'")
            return docs
        except Exception as e:
            print(f"[RAG Retrieval] Error: {e}")
            return []
    
    def _format_context(self, documents: List) -> str:
        """Format retrieved documents into context string"""
        if not documents:
            return "No relevant documents found."
        
        context_parts = []
        for i, doc in enumerate(documents, 1):
            if hasattr(doc, 'metadata'):
                source = doc.metadata.get('filename', 'Unknown')
                content = doc.page_content
            else:
                source = 'Unknown'
                content = str(doc)
            context_parts.append(f"[Chunk {i} from {source}]\n{content}")
        
        return "\n\n".join(context_parts)
    
    async def answer_stream(self, question: str, session_id: str = "default"):
        """
        Answer question using Agentic RAG with streaming
        """
        if not self.client:
            yield "Error: Gemini client not initialized. Please check API key."
            return
        
        # Step 1: Get conversation history
        history_messages = self.get_windowed_history(session_id)
        print(f"[Agentic RAG] Session: {session_id}, History messages: {len(history_messages)}")
        
        # Step 2: Rephrase query using history (History Aware Query Rephraser)
        rephrased_query = await self._rephrase_query(question, history_messages)
        
        # Step 3: Router - Check if web search is needed
        needs_web = self._needs_web_search(question)
        has_documents = self.get_document_count(session_id) > 0
        
        # Step 4: Retrieve relevant context from vector store (if documents exist)
        doc_context = ""
        has_relevant_docs = False
        if has_documents:
            retrieved_docs = self._retrieve_context(rephrased_query, session_id)
            doc_context = self._format_context(retrieved_docs)
            # Check if we actually got relevant content
            has_relevant_docs = doc_context and doc_context != "No relevant documents found."
        
        # Step 5: Search web if needed OR if documents don't have relevant info
        web_context = ""
        if needs_web or (not has_relevant_docs):
            web_context = self._search_web(rephrased_query)
        
        # Combine contexts
        combined_context = ""
        sources_used = []
        
        if doc_context and doc_context != "No relevant documents found.":
            combined_context += f"📄 **From Uploaded Documents:**\n{doc_context}\n\n"
            sources_used.append("documents")
        
        if web_context:
            combined_context += f"🌐 **From Web Search:**\n{web_context}\n\n"
            sources_used.append("web")
        
        if not combined_context:
            combined_context = "No relevant information found in documents or web search."
        
        # Format history for RAG prompt
        history_text = self._format_history_for_rephraser(history_messages)
        
        print(f"[Agentic RAG] Sources used: {sources_used}, Needs web: {needs_web}, Has docs: {has_documents}")
        
        # Step 6: Generate answer using Agentic RAG prompt
        rag_system_prompt = """You are a helpful AI assistant with access to both uploaded documents and web search results.

IMPORTANT RULES:
1. ALWAYS use the provided context to answer - this includes BOTH document context AND web search results
2. If web search results are provided, USE THEM to answer the question directly
3. Do NOT ask for permission to search the web - the search has already been done for you
4. If using documents, reference the source when possible
5. If context includes web results, provide the answer from those results
6. Consider the conversation history for context
7. Use clear, concise language with markdown formatting
8. Do NOT make up information - only use what's provided in the context"""

        rag_user_prompt = f"""Query: {question}

Retrieved Context:
{combined_context}

Conversation History:
{history_text}

Answer the query using the context above. If web search results are provided, use them to give a complete answer. Do NOT ask user permission to search - just use the provided results."""
        
        # Stream the response
        full_response = ""
        try:
            response = self.client.models.generate_content_stream(
                model="gemini-2.0-flash-lite-001",
                contents=f"{rag_system_prompt}\n\n{rag_user_prompt}"
            )
            for chunk in response:
                if chunk.text:
                    full_response += chunk.text
                    yield chunk.text
            
            # Save context after streaming completes
            if full_response:
                self.save_context(session_id, question, full_response)
                
        except Exception as e:
            print(f"[Agentic RAG] Streaming error: {e}")
            yield f"Error generating response: {str(e)}"
    
    async def answer(self, question: str, session_id: str = "default") -> str:
        """Non-streaming version of answer"""
        full_response = ""
        async for chunk in self.answer_stream(question, session_id):
            full_response += chunk
        return full_response
    
    def get_session_documents(self, session_id: str) -> List[str]:
        """Get list of uploaded documents for a session"""
        return self.session_documents.get(session_id, [])
    
    def get_document_count(self, session_id: str) -> int:
        """Get count of documents for a session"""
        return len(self.session_documents.get(session_id, []))


# Global instance
rag_service = RAGService()
